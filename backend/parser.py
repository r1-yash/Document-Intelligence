import io
import pdfplumber
import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

def parse_pdf(file_bytes: bytes) -> str:
    """
    Extracts text page-by-page from a PDF byte stream.
    
    We collect page texts into a list and join them with double newlines
    to preserve visual page separation for the downstream LLM.
    """
    extracted_pages = []
    
    # pdfplumber requires a file-like object, so we wrap bytes in BytesIO
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_index, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                extracted_pages.append(page_text)
                
    return "\n\n".join(extracted_pages)

def parse_docx(file_bytes: bytes) -> str:
    """
    Parses paragraph and table text from a Word document byte stream.
    
    To maintain the document's original logical layout (e.g. paragraphs interspaced 
    with tables in exact order), we iterate directly over the XML body elements 
    rather than accessing `doc.paragraphs` and `doc.tables` independently.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    document_flow = []
    
    try:
        # Loop through elements in the body XML tree to respect page-reading flow
        for element in doc.element.body:
            if isinstance(element, CT_P):
                paragraph = Paragraph(element, doc)
                if paragraph.text.strip():
                    document_flow.append(paragraph.text)
                    
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                formatted_table_rows = []
                
                # Word table columns may contain merged cells. In the underlying python-docx 
                # API, accessing row.cells yields duplicate references for merged fields.
                # We track seen cells to prevent printing the same content multiple times.
                for row in table.rows:
                    row_data = []
                    seen_cells = set()
                    
                    for cell in row.cells:
                        if cell not in seen_cells:
                            seen_cells.add(cell)
                            # Clean up inline breaks to keep table rows structured as single lines
                            cleaned_cell_text = cell.text.replace("\n", " ").strip()
                            row_data.append(cleaned_cell_text)
                            
                    formatted_table_rows.append(" | ".join(row_data))
                    
                if formatted_table_rows:
                    document_flow.append("\n".join(formatted_table_rows))
                    
    except Exception:
        # Safeguard: if parsing via XML tree elements fails (due to schema drift or unusual formats),
        # fall back to standard sequencial dumps of paragraphs and then tables.
        document_flow = [p.text for p in doc.paragraphs if p.text.strip()]
        
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                seen_cells = set()
                for cell in row.cells:
                    if cell not in seen_cells:
                        seen_cells.add(cell)
                        row_data.append(cell.text.replace("\n", " ").strip())
                document_flow.append(" | ".join(row_data))
                
    return "\n\n".join(document_flow)

def parse_txt(file_bytes: bytes) -> str:
    """
    Attempts to decode raw text bytes using common standard encodings.
    
    Many corporate TXT dumps originate from older systems with ISO-8859-1 (Latin-1) 
    or UTF-16 schemas, so we cycle through fallbacks if UTF-8 raises a decoding error.
    """
    target_encodings = ["utf-8", "latin-1", "utf-16"]
    
    for encoding in target_encodings:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
            
    raise ValueError(
        "Could not decode text file. Ensure it is encoded in UTF-8, Latin-1, or UTF-16."
    )

def parse_document(filename: str, file_bytes: bytes) -> str:
    """
    Dispatches a document to its corresponding parser based on file extension.
    """
    file_extension = filename.split(".")[-1].lower()
    
    if file_extension == "pdf":
        return parse_pdf(file_bytes)
    elif file_extension in ("docx", "doc"):
        return parse_docx(file_bytes)
    elif file_extension == "txt":
        return parse_txt(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file format: '.{file_extension}'. Supported types are PDF, DOCX, and TXT."
        )
